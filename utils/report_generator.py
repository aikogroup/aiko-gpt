"""
Générateur de rapport Word (.docx) pour les résultats d'analyse IA
"""

import os
from datetime import datetime
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement
import config


class ReportGenerator:
    """
    Générateur de rapports Word pour les résultats d'analyse des besoins et cas d'usage IA
    """
    
    def __init__(self, logo_path: str = None):
        """
        Initialise le générateur de rapport.
        
        Args:
            logo_path: Chemin vers le logo Aiko (optionnel)
        """
        self.logo_path = logo_path
        if not logo_path:
            # Utiliser le chemin depuis config.py (détection automatique)
            self.logo_path = str(config.get_logo_path())

    def _remove_numbering_from_paragraph(self, paragraph):
        """
        Supprime un éventuel <w:numPr> (numérotation automatique) du paragraphe
        pour éviter les numéros automatiques doublés.
        """
        try:
            # Accéder directement à l'élément XML du paragraphe
            p_element = paragraph._p
            
            # Utiliser le qualified name pour le tag numPr
            num_pr_qname = qn('w:numPr')
            
            # Parcourir tous les éléments du paragraphe et supprimer les numPr trouvés
            # Cette approche évite l'utilisation de xpath avec namespaces qui peut poser problème
            elements_to_remove = []
            for elem in p_element.iter():
                if elem.tag == num_pr_qname:
                    elements_to_remove.append(elem)
            
            # Supprimer les éléments trouvés
            for num_pr in elements_to_remove:
                parent = num_pr.getparent()
                if parent is not None:
                    parent.remove(num_pr)
        except Exception as e:
            # Si une erreur survient, on continue sans supprimer la numérotation
            # pour éviter de bloquer la génération du rapport
            print(f"⚠️ [REPORT] Erreur lors de la suppression de la numérotation : {str(e)}")
    
    def generate_report(
        self,
        company_name: str,
        final_needs: List[Dict[str, Any]],
        final_use_cases: List[Dict[str, Any]],
        output_dir: str = None
    ) -> str:
        """
        Génère un rapport Word complet selon le template Cousin Surgery.
        
        Args:
            company_name: Nom de l'entreprise
            final_needs: Liste des besoins identifiés
            final_use_cases: Liste des cas d'usage IA
            output_dir: Dossier de sortie
            
        Returns:
            Chemin vers le fichier généré
        """
        print(f"📝 [REPORT] Génération du rapport pour {company_name}")
        
        # Utiliser le dossier de sortie par défaut depuis config.py si non spécifié
        if output_dir is None:
            output_dir = str(config.ensure_outputs_dir())
        
        # Formater le nom de l'entreprise (première lettre de chaque mot en majuscule)
        company_name_formatted = company_name.title() if company_name else company_name
        print(f"✨ [REPORT] Nom formaté: {company_name_formatted}")
        
        # Créer un nouveau document
        doc = Document()
        
        # Configuration du document
        self._setup_document_styles(doc)
        
        # Ajouter le logo Aiko si disponible
        if os.path.exists(self.logo_path):
            self._add_logo(doc)
        else:
            print(f"⚠️ [REPORT] Logo Aiko non trouvé : {self.logo_path}")
        
        # Ajouter le contenu (utiliser le nom formaté)
        self._add_needs_section(doc, company_name_formatted, final_needs)
        self._add_use_cases_section(doc, company_name_formatted, final_use_cases)
        
        # Générer le nom du fichier (utiliser le nom formaté)
        date_str = datetime.now().strftime("%d%m")
        filename = f"{date_str}-V0-Cas_d_usages_IA-{company_name_formatted.replace(' ', '_')}.docx"
        output_path = os.path.join(output_dir, filename)
        
        # Sauvegarder le document
        doc.save(output_path)
        
        print(f"✅ [REPORT] Rapport généré : {output_path}")
        return output_path
    
    def _setup_document_styles(self, doc: Document):
        """
        Configure les styles du document pour matcher le template Cousin Surgery.
        
        Args:
            doc: Document docx
        """
        # Configuration des marges
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        # Modifier les styles par défaut
        styles = doc.styles
        
        # Style pour Heading 2
        if 'Heading 2' in styles:
            h2_style = styles['Heading 2']
            h2_font = h2_style.font
            h2_font.size = Pt(16)
            h2_font.bold = True
            h2_font.color.rgb = RGBColor(31, 73, 125)  # Bleu foncé
        
        # Style pour Heading 4
        if 'Heading 4' in styles:
            h4_style = styles['Heading 4']
            h4_font = h4_style.font
            h4_font.size = Pt(12)
            h4_font.bold = True
            h4_font.color.rgb = RGBColor(31, 73, 125)
        
        # Créer un style personnalisé pour les titres de famille
        self._create_family_heading_style(doc)
    
    def _create_family_heading_style(self, doc: Document):
        """
        Crée un style personnalisé pour les titres de famille des use cases.
        
        Args:
            doc: Document docx
        """
        styles = doc.styles
        
        # Créer un style personnalisé si il n'existe pas
        style_name = "FamilyHeading"
        if style_name not in styles:
            # Créer un nouveau style basé sur Normal
            family_style = styles.add_style(style_name, 1)  # 1 = paragraph style
        else:
            family_style = styles[style_name]
        
        # Configurer la police
        font = family_style.font
        font.name = "DM Sans"
        font.size = Pt(14)
        font.bold = True
        font.color.rgb = RGBColor(31, 73, 125)  # Même couleur que "LES BESOINS IDENTIFIÉS"
        
        # Configurer le paragraphe (sans puce, espacement)
        paragraph_format = family_style.paragraph_format
        paragraph_format.space_before = Pt(12)
        paragraph_format.space_after = Pt(6)
        paragraph_format.left_indent = Inches(0)
    
    def _add_logo(self, doc: Document):
        """
        Ajoute le logo Aiko en haut du document.
        
        Args:
            doc: Document docx
        """
        try:
            # Ajouter le logo au début du document
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run()
            run.add_picture(self.logo_path, width=Inches(1.5))
            
            # Ajouter un espace après le logo
            doc.add_paragraph()
            
            print(f"✅ [REPORT] Logo Aiko ajouté")
        except Exception as e:
            print(f"❌ [REPORT] Erreur lors de l'ajout du logo : {str(e)}")
    
    def _add_needs_section(self, doc: Document, company_name: str, needs: List[Dict[str, Any]]):
        """
        Ajoute la section des besoins identifiés.
        
        Args:
            doc: Document docx
            company_name: Nom de l'entreprise
            needs: Liste des besoins
        """
        # Titre principal
        heading = doc.add_heading(f"LES BESOINS IDENTIFIÉS DE {company_name.upper()}", level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Organiser les besoins par thème
        needs_by_theme = {}
        for need in needs:
            theme = need.get('theme', 'Autres besoins')
            if theme not in needs_by_theme:
                needs_by_theme[theme] = []
            needs_by_theme[theme].append(need)
        
        # Ajouter chaque thème
        for theme, theme_needs in needs_by_theme.items():
            # Sous-titre avec emoji
            subheading = doc.add_heading(f"🔹 {theme}", level=4)
            
            # Ajouter les citations
            for need in theme_needs:
                quotes = need.get('quotes', [])
                for quote in quotes:
                    # Nettoyer les guillemets si nécessaire
                    cleaned_quote = quote.strip()
                    if not cleaned_quote.startswith('«'):
                        cleaned_quote = f"« {cleaned_quote}"
                    if not cleaned_quote.endswith('»'):
                        cleaned_quote = f"{cleaned_quote} »"
                    
                    # OPTION B (robuste): puce manuelle + indentation
                    paragraph = doc.add_paragraph(style='Normal')
                    paragraph.paragraph_format.left_indent = Inches(0.6)
                    paragraph.paragraph_format.space_after = Pt(6)
                    # ajouter la puce manuelle pour éviter les comportements de style Word
                    run = paragraph.add_run("• ")
                    run.bold = False
                    run_quote = paragraph.add_run(cleaned_quote)

    
    def _add_use_cases_section(
        self,
        doc: Document,
        company_name: str,
        use_cases: List[Dict[str, Any]]
    ):
        """
        Ajoute la section des cas d'usage IA.
        
        Args:
            doc: Document docx
            company_name: Nom de l'entreprise
            use_cases: Liste des cas d'usage IA
        """
        # Titre principal
        doc.add_heading("LES CAS D'USAGES IA PRIORITAIRES", level=2)
        
        # Paragraphe d'introduction
        intro_text = (
            f"Pour donner suite à la série d'entretiens et aux ateliers de travail menés avec "
            f"les équipes de {company_name}, nous avons identifié des cas d'usage qui répondent "
            f"directement aux besoins et enjeux stratégiques IA de l'entreprise. "
            f"Voici les cas d'usage prioritaires qui émergent de cette réflexion collective :"
        )
        doc.add_paragraph(intro_text)
        
        # Grouper les use cases par famille
        if use_cases:
            use_cases_by_family = {}
            use_cases_without_family = []
            
            for uc in use_cases:
                famille = uc.get('famille', '').strip() if uc.get('famille') else None
                if famille:
                    if famille not in use_cases_by_family:
                        use_cases_by_family[famille] = []
                    use_cases_by_family[famille].append(uc)
                else:
                    use_cases_without_family.append(uc)
            
            # Compteur global pour la numérotation continue
            global_counter = 1
            
            # Afficher les use cases groupés par famille
            for famille, famille_use_cases in use_cases_by_family.items():
                # Titre de famille avec style personnalisé
                family_para = doc.add_paragraph(style='FamilyHeading')
                family_run = family_para.add_run(famille)
                family_run.bold = True
                family_run.font.name = "DM Sans"
                family_run.font.size = Pt(14)
                family_run.font.color.rgb = RGBColor(31, 73, 125)
                
                # Afficher les use cases de cette famille
                for uc in famille_use_cases:
                    title = uc.get('titre', 'N/A')
                    description = uc.get('description', 'N/A')
                    
                    # Titre : on force style Normal pour éviter héritage
                    title_para = doc.add_paragraph(style='Normal')
                    run_title = title_para.add_run(f"{global_counter}. {title}")
                    run_title.bold = True

                    # Retirer toute numérotation automatique cachée
                    self._remove_numbering_from_paragraph(title_para)

                    # Description : run 1 en gras pour "Description : ", run 2 pour le texte
                    desc_para = doc.add_paragraph(style='Normal')
                    desc_para.paragraph_format.left_indent = Inches(0.4)
                    desc_para.paragraph_format.space_after = Pt(12)

                    run_desc_label = desc_para.add_run("Description : ")
                    run_desc_label.bold = True
                    run_desc = desc_para.add_run(description)
                    
                    global_counter += 1
            
            # Afficher les use cases sans famille dans une section "Autres cas d'usage"
            if use_cases_without_family:
                # Titre de section "Autres cas d'usage"
                others_para = doc.add_paragraph(style='FamilyHeading')
                others_run = others_para.add_run("Autres cas d'usage")
                others_run.bold = True
                others_run.font.name = "DM Sans"
                others_run.font.size = Pt(14)
                others_run.font.color.rgb = RGBColor(31, 73, 125)
                
                # Afficher les use cases sans famille
                for uc in use_cases_without_family:
                    title = uc.get('titre', 'N/A')
                    description = uc.get('description', 'N/A')
                    
                    # Titre : on force style Normal pour éviter héritage
                    title_para = doc.add_paragraph(style='Normal')
                    run_title = title_para.add_run(f"{global_counter}. {title}")
                    run_title.bold = True

                    # Retirer toute numérotation automatique cachée
                    self._remove_numbering_from_paragraph(title_para)

                    # Description : run 1 en gras pour "Description : ", run 2 pour le texte
                    desc_para = doc.add_paragraph(style='Normal')
                    desc_para.paragraph_format.left_indent = Inches(0.4)
                    desc_para.paragraph_format.space_after = Pt(12)

                    run_desc_label = desc_para.add_run("Description : ")
                    run_desc_label.bold = True
                    run_desc = desc_para.add_run(description)
                    
                    global_counter += 1
    
    def generate_report_from_json_files(
        self,
        company_name: str,
        needs_json_path: str = None,
        use_cases_json_path: str = None,
        output_dir: str = None
    ) -> str:
        """
        Génère un rapport à partir des fichiers JSON sauvegardés.
        
        Args:
            company_name: Nom de l'entreprise
            needs_json_path: Chemin vers le JSON des besoins
            use_cases_json_path: Chemin vers le JSON des cas d'usage
            output_dir: Dossier de sortie
            
        Returns:
            Chemin vers le fichier généré
        """
        import json
        
        # Utiliser les chemins par défaut depuis config.py si non spécifiés
        if needs_json_path is None:
            needs_json_path = str(config.OUTPUTS_DIR / "need_analysis_results.json")
        if use_cases_json_path is None:
            use_cases_json_path = str(config.OUTPUTS_DIR / "use_case_analysis_results.json")
        if output_dir is None:
            output_dir = str(config.ensure_outputs_dir())
        
        print(f"📂 [REPORT] Chargement des données depuis les fichiers JSON")
        
        # Charger les besoins
        try:
            with open(needs_json_path, 'r', encoding='utf-8') as f:
                needs_data = json.load(f)
            final_needs = needs_data.get('final_needs', [])
            print(f"✅ [REPORT] {len(final_needs)} besoins chargés")
        except FileNotFoundError:
            print(f"⚠️ [REPORT] Fichier des besoins non trouvé : {needs_json_path}")
            final_needs = []
        except Exception as e:
            print(f"❌ [REPORT] Erreur lors du chargement des besoins : {str(e)}")
            final_needs = []
        
        # Charger les cas d'usage
        try:
            with open(use_cases_json_path, 'r', encoding='utf-8') as f:
                use_cases_data = json.load(f)
            # Tenter de charger final_use_cases (nouveau format)
            final_use_cases = use_cases_data.get('final_use_cases', [])
            # Si absent, essayer l'ancien format pour rétrocompatibilité
            if not final_use_cases:
                final_quick_wins = use_cases_data.get('final_quick_wins', [])
                final_structuration_ia = use_cases_data.get('final_structuration_ia', [])
                final_use_cases = final_quick_wins + final_structuration_ia
                print(f"✅ [REPORT] {len(final_use_cases)} cas d'usage chargés (ancien format)")
            else:
                print(f"✅ [REPORT] {len(final_use_cases)} cas d'usage chargés")
        except FileNotFoundError:
            print(f"⚠️ [REPORT] Fichier des cas d'usage non trouvé : {use_cases_json_path}")
            final_use_cases = []
        except Exception as e:
            print(f"❌ [REPORT] Erreur lors du chargement des cas d'usage : {str(e)}")
            final_use_cases = []
        
        # Générer le rapport
        return self.generate_report(
            company_name=company_name,
            final_needs=final_needs,
            final_use_cases=final_use_cases,
            output_dir=output_dir
        )

