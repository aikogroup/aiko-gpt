"""
Générateur de rapport Word (.docx) pour les résultats d'analyse IA
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class ReportGenerator:
    """
    Générateur de rapports Word pour les résultats d'analyse des besoins et cas d'usage IA
    """
    
    def __init__(self, logo_path: str = None):
        """
        Initialise le générateur de rapport.
        
        Args:
            logo_path: Chemin vers le logo Aiko (optionnel)
        """
        self.logo_path = logo_path
        if not logo_path:
            # Chemin par défaut pour le logo Aiko (mis à jour pour utiliser le logo frontend s'il est synchronisé côté API)
            # On tente d'abord dans un dossier public standard, sinon fallback sur l'ancien chemin si présent
            candidate_paths = [
                str(Path(__file__).parent.parent / "frontend" / "public" / "logoAiko.jpeg"),
                "/home/addeche/aiko/aikoGPT/assets/aiko_logo.png",
            ]
            for p in candidate_paths:
                if os.path.exists(p):
                    self.logo_path = p
                    break
    
    def generate_report(
        self,
        company_name: str,
        final_needs: List[Dict[str, Any]],
        final_quick_wins: List[Dict[str, Any]],
        final_structuration_ia: List[Dict[str, Any]],
        output_dir: str = "/home/addeche/aiko/aikoGPT/outputs"
    ) -> str:
        """
        Génère un rapport Word complet selon le template Cousin Surgery.
        
        Args:
            company_name: Nom de l'entreprise
            final_needs: Liste des besoins identifiés
            final_quick_wins: Liste des cas d'usage Quick Wins
            final_structuration_ia: Liste des cas d'usage Structuration IA
            output_dir: Dossier de sortie
            
        Returns:
            Chemin vers le fichier généré
        """
        print(f"📝 [REPORT] Génération du rapport pour {company_name}")
        
        # Formater le nom de l'entreprise (première lettre de chaque mot en majuscule)
        company_name_formatted = company_name.title() if company_name else company_name
        print(f"✨ [REPORT] Nom formaté: {company_name_formatted}")
        
        # Créer un nouveau document
        doc = Document()
        
        # Configuration du document
        self._setup_document_styles(doc)
        
        # Ajouter le logo Aiko si disponible
        if os.path.exists(self.logo_path):
            self._add_logo(doc)
        else:
            print(f"⚠️ [REPORT] Logo Aiko non trouvé : {self.logo_path}")
        
        # Ajouter le contenu (utiliser le nom formaté)
        self._add_needs_section(doc, company_name_formatted, final_needs)
        self._add_use_cases_section(doc, company_name_formatted, final_quick_wins, final_structuration_ia)
        
        # Générer le nom du fichier (utiliser le nom formaté)
        date_str = datetime.now().strftime("%d%m")
        filename = f"{date_str}-V0-Cas_d_usages_IA-{company_name_formatted.replace(' ', '_')}.docx"
        output_path = os.path.join(output_dir, filename)
        
        # Sauvegarder le document
        doc.save(output_path)
        
        print(f"✅ [REPORT] Rapport généré : {output_path}")
        return output_path
    
    def _setup_document_styles(self, doc: Document):
        """
        Configure les styles du document pour matcher le template Cousin Surgery.
        
        Args:
            doc: Document docx
        """
        # Configuration des marges
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        # Modifier les styles par défaut
        styles = doc.styles
        
        # Style pour Heading 2
        if 'Heading 2' in styles:
            h2_style = styles['Heading 2']
            h2_font = h2_style.font
            h2_font.size = Pt(16)
            h2_font.bold = True
            h2_font.color.rgb = RGBColor(31, 73, 125)  # Bleu foncé
        
        # Style pour Heading 4
        if 'Heading 4' in styles:
            h4_style = styles['Heading 4']
            h4_font = h4_style.font
            h4_font.size = Pt(12)
            h4_font.bold = True
            h4_font.color.rgb = RGBColor(31, 73, 125)
    
    def _add_logo(self, doc: Document):
        """
        Ajoute le logo Aiko en haut du document.
        
        Args:
            doc: Document docx
        """
        try:
            # Ajouter le logo au début du document
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run()
            run.add_picture(self.logo_path, width=Inches(1.5))
            
            # Ajouter un espace après le logo
            doc.add_paragraph()
            
            print(f"✅ [REPORT] Logo Aiko ajouté")
        except Exception as e:
            print(f"❌ [REPORT] Erreur lors de l'ajout du logo : {str(e)}")
    
    def _add_needs_section(self, doc: Document, company_name: str, needs: List[Dict[str, Any]]):
        """
        Ajoute la section des besoins identifiés.
        
        Args:
            doc: Document docx
            company_name: Nom de l'entreprise
            needs: Liste des besoins
        """
        # Titre principal
        heading = doc.add_heading(f"LES BESOINS IDENTIFIÉS DE {company_name.upper()}", level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Organiser les besoins par thème
        needs_by_theme = {}
        for need in needs:
            theme = need.get('theme', 'Autres besoins')
            if theme not in needs_by_theme:
                needs_by_theme[theme] = []
            needs_by_theme[theme].append(need)
        
        # Ajouter chaque thème
        for theme, theme_needs in needs_by_theme.items():
            # Sous-titre avec emoji
            subheading = doc.add_heading(f"🔹 {theme}", level=4)
            
            # Ajouter les citations
            for need in theme_needs:
                quotes = need.get('quotes', [])
                for quote in quotes:
                    # Nettoyer les guillemets si nécessaire
                    cleaned_quote = quote.strip()
                    if not cleaned_quote.startswith('«'):
                        cleaned_quote = f"« {cleaned_quote}"
                    if not cleaned_quote.endswith('»'):
                        cleaned_quote = f"{cleaned_quote} »"
                    
                    # Ajouter comme élément de liste
                    paragraph = doc.add_paragraph(cleaned_quote, style='List Paragraph')
    
    def _add_use_cases_section(
        self,
        doc: Document,
        company_name: str,
        quick_wins: List[Dict[str, Any]],
        structuration_ia: List[Dict[str, Any]]
    ):
        """
        Ajoute la section des cas d'usage IA.
        
        Args:
            doc: Document docx
            company_name: Nom de l'entreprise
            quick_wins: Liste des Quick Wins
            structuration_ia: Liste des cas Structuration IA
        """
        # Titre principal
        doc.add_heading("LES CAS D'USAGES IA PRIORITAIRES", level=2)
        
        # Paragraphe d'introduction
        intro_text = (
            f"Pour donner suite à la série d'entretiens et aux ateliers de travail menés avec "
            f"les équipes de {company_name}, nous avons identifié des cas d'usage qui répondent "
            f"directement aux besoins et enjeux stratégiques IA de l'entreprise. "
            f"Voici les cas d'usage prioritaires qui émergent de cette réflexion collective :"
        )
        doc.add_paragraph(intro_text)
        
        # Section Quick Wins
        if quick_wins:
            doc.add_heading('Famille "Quick Wins" – Automatisation & assistance intelligente', level=2)
            
            for uc in quick_wins:
                # Titre du cas d'usage
                paragraph = doc.add_paragraph(uc.get('titre', 'N/A'), style='List Paragraph')
                
                # IA utilisée
                ia_para = doc.add_paragraph(style='List Paragraph')
                ia_run = ia_para.add_run(f"IA : {uc.get('ia_utilisee', 'N/A')}")
                ia_run.italic = True
                
                # Description
                desc_para = doc.add_paragraph(style='List Paragraph')
                desc_run = desc_para.add_run(f"Description : {uc.get('description', 'N/A')}")
        
        # Section Structuration IA
        if structuration_ia:
            doc.add_heading(
                'Famille "Structuration IA à moyen et long terme" Scalabilité & qualité prédictive',
                level=2
            )
            
            for uc in structuration_ia:
                # Titre du cas d'usage
                paragraph = doc.add_paragraph(uc.get('titre', 'N/A'), style='List Paragraph')
                
                # IA utilisée
                ia_para = doc.add_paragraph(style='List Paragraph')
                ia_run = ia_para.add_run(f"IA : {uc.get('ia_utilisee', 'N/A')}")
                ia_run.italic = True
                
                # Description
                desc_para = doc.add_paragraph(style='List Paragraph')
                desc_run = desc_para.add_run(f"Description : {uc.get('description', 'N/A')}")
    
    def generate_report_from_json_files(
        self,
        company_name: str,
        needs_json_path: str = "/home/addeche/aiko/aikoGPT/outputs/need_analysis_results.json",
        use_cases_json_path: str = "/home/addeche/aiko/aikoGPT/outputs/use_case_analysis_results.json",
        output_dir: str = "/home/addeche/aiko/aikoGPT/outputs"
    ) -> str:
        """
        Génère un rapport à partir des fichiers JSON sauvegardés.
        
        Args:
            company_name: Nom de l'entreprise
            needs_json_path: Chemin vers le JSON des besoins
            use_cases_json_path: Chemin vers le JSON des cas d'usage
            output_dir: Dossier de sortie
            
        Returns:
            Chemin vers le fichier généré
        """
        import json
        
        print(f"📂 [REPORT] Chargement des données depuis les fichiers JSON")
        
        # Charger les besoins
        try:
            with open(needs_json_path, 'r', encoding='utf-8') as f:
                needs_data = json.load(f)
            final_needs = needs_data.get('final_needs', [])
            print(f"✅ [REPORT] {len(final_needs)} besoins chargés")
        except FileNotFoundError:
            print(f"⚠️ [REPORT] Fichier des besoins non trouvé : {needs_json_path}")
            final_needs = []
        except Exception as e:
            print(f"❌ [REPORT] Erreur lors du chargement des besoins : {str(e)}")
            final_needs = []
        
        # Charger les cas d'usage
        try:
            with open(use_cases_json_path, 'r', encoding='utf-8') as f:
                use_cases_data = json.load(f)
            final_quick_wins = use_cases_data.get('final_quick_wins', [])
            final_structuration_ia = use_cases_data.get('final_structuration_ia', [])
            print(f"✅ [REPORT] {len(final_quick_wins)} Quick Wins et {len(final_structuration_ia)} Structuration IA chargés")
        except FileNotFoundError:
            print(f"⚠️ [REPORT] Fichier des cas d'usage non trouvé : {use_cases_json_path}")
            final_quick_wins = []
            final_structuration_ia = []
        except Exception as e:
            print(f"❌ [REPORT] Erreur lors du chargement des cas d'usage : {str(e)}")
            final_quick_wins = []
            final_structuration_ia = []
        
        # Générer le rapport
        return self.generate_report(
            company_name=company_name,
            final_needs=final_needs,
            final_quick_wins=final_quick_wins,
            final_structuration_ia=final_structuration_ia,
            output_dir=output_dir
        )

