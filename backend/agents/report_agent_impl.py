"""
ReportAgent Implementation - Génération du rapport Word final

FR: Implémentation complète du ReportAgent avec génération document Word professionnel
"""

import logging
from typing import Dict, Any, List
from pathlib import Path
from datetime import datetime

# FR: Import optionnel de python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

from models.graph_state import NeedAnalysisState

logger = logging.getLogger(__name__)


def create_word_report(
    validated_needs: List[Dict[str, Any]],
    validated_use_cases: List[Dict[str, Any]],
    company_name: str,
    output_path: str
) -> str:
    """
    FR: Crée le document Word avec besoins et cas d'usage
    
    Args:
        validated_needs: Besoins sélectionnés
        validated_use_cases: Cas d'usage sélectionnés (QW + SIA)
        company_name: Nom de l'entreprise
        output_path: Chemin de sortie du fichier
        
    Returns:
        str: Chemin du fichier généré
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx n'est pas installé - génération Word impossible")
    
    logger.info(f"📝 Création du document Word: {output_path}")
    
    # FR: Créer le document
    doc = Document()
    
    # FR: Style du titre principal
    title = doc.add_heading(f"Analyse des Besoins IA - {company_name}", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # FR: Date du rapport
    date_para = doc.add_paragraph()
    date_para.add_run(f"Rapport généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}").italic = True
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_page_break()
    
    # FR: Section 1 - Besoins validés
    doc.add_heading("1. Besoins Métier Identifiés", 1)
    doc.add_paragraph(
        f"Cette section présente les {len(validated_needs)} besoins métier prioritaires "
        f"identifiés lors de l'analyse des ateliers et des entretiens collaborateurs."
    )
    doc.add_paragraph()
    
    for idx, need in enumerate(validated_needs, 1):
        # FR: Titre du besoin
        need_heading = doc.add_heading(f"{idx}. {need.get('title', 'Besoin sans titre')}", 2)
        
        # FR: Citations associées
        doc.add_paragraph("Citations issues des ateliers et entretiens :", style='Intense Quote')
        for citation in need.get('citations', [])[:5]:
            p = doc.add_paragraph(citation, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # FR: Section 2 - Cas d'usage IA
    doc.add_heading("2. Cas d'Usage IA Proposés", 1)
    doc.add_paragraph(
        f"Cette section présente les {len(validated_use_cases)} cas d'usage IA retenus, "
        f"organisés par type (Quick Wins et Structuration IA)."
    )
    doc.add_paragraph()
    
    # FR: Séparer Quick Wins et Structuration IA
    quick_wins = [uc for uc in validated_use_cases if uc.get('category') == 'quick_win']
    structuration_ia = [uc for uc in validated_use_cases if uc.get('category') == 'structuration_ia']
    
    # FR: Quick Wins
    if quick_wins:
        doc.add_heading("2.1 Quick Wins (ROI immédiat < 3 mois)", 2)
        doc.add_paragraph(
            "Solutions à faible complexité technique et mise en œuvre rapide."
        )
        doc.add_paragraph()
        
        for idx, uc in enumerate(quick_wins, 1):
            # FR: Titre du cas d'usage
            uc_heading = doc.add_paragraph()
            title = uc.get('title', "Cas d'usage sans titre")
            uc_heading.add_run(f"QW{idx}. {title}").bold = True
            uc_heading.paragraph_format.left_indent = Inches(0.25)
            
            # FR: Description
            desc_para = doc.add_paragraph(uc.get('description', 'Description non disponible'))
            desc_para.paragraph_format.left_indent = Inches(0.5)
            
            # FR: Technologies IA
            tech_para = doc.add_paragraph()
            tech_para.add_run("Technologies IA : ").bold = True
            tech_para.add_run(", ".join(uc.get('ai_technologies', [])))
            tech_para.paragraph_format.left_indent = Inches(0.5)
            
            doc.add_paragraph()
    
    # FR: Structuration IA
    if structuration_ia:
        doc.add_heading("2.2 Structuration IA (ROI moyen/long terme 3-12 mois)", 2)
        doc.add_paragraph(
            "Solutions avancées avec complexité moyenne/élevée et mise en œuvre progressive."
        )
        doc.add_paragraph()
        
        for idx, uc in enumerate(structuration_ia, 1):
            # FR: Titre du cas d'usage
            uc_heading = doc.add_paragraph()
            title = uc.get('title', "Cas d'usage sans titre")
            uc_heading.add_run(f"SIA{idx}. {title}").bold = True
            uc_heading.paragraph_format.left_indent = Inches(0.25)
            
            # FR: Description
            desc_para = doc.add_paragraph(uc.get('description', 'Description non disponible'))
            desc_para.paragraph_format.left_indent = Inches(0.5)
            
            # FR: Technologies IA
            tech_para = doc.add_paragraph()
            tech_para.add_run("Technologies IA : ").bold = True
            tech_para.add_run(", ".join(uc.get('ai_technologies', [])))
            tech_para.paragraph_format.left_indent = Inches(0.5)
            
            doc.add_paragraph()
    
    # FR: Section 3 - Résumé
    doc.add_page_break()
    doc.add_heading("3. Résumé Exécutif", 1)
    
    summary_para = doc.add_paragraph()
    summary_para.add_run(f"Besoins identifiés : ").bold = True
    summary_para.add_run(f"{len(validated_needs)}\n")
    summary_para.add_run(f"Quick Wins proposés : ").bold = True
    summary_para.add_run(f"{len(quick_wins)}\n")
    summary_para.add_run(f"Structuration IA proposée : ").bold = True
    summary_para.add_run(f"{len(structuration_ia)}\n")
    summary_para.add_run(f"Total cas d'usage : ").bold = True
    summary_para.add_run(f"{len(validated_use_cases)}")
    
    # FR: Sauvegarder le document
    doc.save(output_path)
    logger.info(f"✅ Document Word sauvegardé: {output_path}")
    
    return output_path


def report_agent(state: NeedAnalysisState, config: Dict[str, Any]) -> Dict[str, Any]:
    """
    FR: Agent LangGraph pour générer le rapport Word final
    
    Args:
        state: État actuel du workflow LangGraph
        config: Configuration LangGraph
        
    Returns:
        Dict: Mise à jour de l'état avec report_path
    """
    logger.info("📄 ReportAgent - Début génération rapport Word")
    
    # FR: Récupérer les données validées
    validated_needs = state.get("validated_needs", [])
    validated_quick_wins = state.get("validated_quick_wins", [])
    validated_structuration_ia = state.get("validated_structuration_ia", [])
    company_name = state.get("company_name", "Entreprise")
    
    # FR: Combiner tous les cas d'usage validés
    validated_use_cases = validated_quick_wins + validated_structuration_ia
    
    # FR: Vérifier qu'il y a des données à exporter
    if not validated_needs and not validated_use_cases:
        logger.warning("⚠️ Aucun besoin ou cas d'usage validé - rapport vide")
        return {
            "report_path": None,
            "current_step": "report_empty",
            "errors": ["Aucun besoin ou cas d'usage à exporter"]
        }
    
    try:
        if not DOCX_AVAILABLE:
            logger.error("❌ python-docx n'est pas installé")
            return {
                "report_path": None,
                "current_step": "report_error",
                "errors": ["python-docx non installé - génération Word impossible"]
            }
        
        # FR: Créer le dossier de sortie si nécessaire
        output_dir = Path(config.get("output_dir", "./outputs"))
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # FR: Nom du fichier avec timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Rapport_Besoins_IA_{company_name.replace(' ', '_')}_{timestamp}.docx"
        output_path = str(output_dir / filename)
        
        # FR: Générer le document Word
        report_path = create_word_report(
            validated_needs,
            validated_use_cases,
            company_name,
            output_path
        )
        
        logger.info("✅ ReportAgent - Rapport généré")
        logger.info(f"📊 {len(validated_needs)} besoins + {len(validated_use_cases)} cas d'usage")
        
        return {
            "report_path": report_path,
            "current_step": "report_completed"
        }
        
    except Exception as e:
        logger.error(f"❌ Erreur inattendue dans ReportAgent: {e}")
        import traceback
        traceback.print_exc()
        return {
            "report_path": None,
            "current_step": "report_error",
            "errors": [f"Erreur ReportAgent: {str(e)}"]
        }

