"""
Extracteur de données depuis un rapport Word généré
"""

import logging
from typing import Dict, List, Any, Optional
from pathlib import Path
from docx import Document
from openai import OpenAI
import os
from dotenv import load_dotenv
import json
import re
from models.executive_summary_models import WordReportExtraction
from prompts.executive_summary_prompts import WORD_REPORT_EXTRACTION_PROMPT

load_dotenv()

logger = logging.getLogger(__name__)


class WordReportExtractor:
    """Extracteur de données depuis un rapport Word"""
    
    def __init__(self, api_key: str = None):
        """
        Initialise l'extracteur.
        
        Args:
            api_key: Clé API OpenAI (optionnel, utilise OPENAI_API_KEY par défaut)
        """
        api_key = api_key or os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY doit être définie")
        
        self.client = OpenAI(api_key=api_key)
        self.model = os.getenv('OPENAI_MODEL', 'gpt-5-nano')
    
    def extract_from_word(self, word_path: str) -> Dict[str, List[Dict[str, Any]]]:
        """
        Extrait les données depuis un fichier Word.
        
        Stratégie :
        1. Tentative d'extraction JSON directe (métadonnées ou structure cachée)
        2. Si échec : extraction via LLM avec structured output
        
        Args:
            word_path: Chemin vers le fichier Word (.docx)
            
        Returns:
            Dict avec keys: 'final_needs', 'final_use_cases'
        """
        logger.info(f"Extraction depuis Word: {word_path}")
        
        # Étape 1 : Tentative d'extraction structurée (parsing du document)
        extracted_data = self._try_extract_json(word_path)
        if extracted_data:
            logger.info("✅ Extraction structurée réussie (parsing direct)")
            return extracted_data
        
        # Étape 2 : Extraction via LLM (si le parsing échoue)
        logger.info("⚠️ Extraction structurée échouée, utilisation LLM")
        return self._extract_with_llm(word_path)
    
    def _try_extract_json(self, word_path: str) -> Optional[Dict[str, List[Dict[str, Any]]]]:
        """
        Parse le document Word pour extraire les données structurées.
        
        Cette méthode parse le document Word et extrait les besoins et cas d'usage
        basés sur la structure du document :
        - Besoins : lignes commençant par 🔹, avec citations entre « et »
        - Cas d'usage : section "LES CAS D'USAGES IA PRIORITAIRES"
          La famille (si présente) est extraite des titres de famille (style FamilyHeading)
        
        Si l'extraction échoue ou ne trouve rien, on utilise l'extraction LLM.
        
        Args:
            word_path: Chemin vers le fichier Word
            
        Returns:
            Dict avec 'final_needs', 'final_use_cases'
            ou None si échec (on utilisera alors LLM)
        """
        try:
            doc = Document(word_path)
            
            needs = []
            use_cases = []
            
            current_need = None
            current_section = "needs"  # "needs" ou "use_cases"
            current_use_case = None
            current_family = None  # Famille courante pour les use cases
            
            for para in doc.paragraphs:
                text = para.text.strip()
                
                if not text:
                    continue
                
                # --- Détection de section ---
                if "LES CAS D'USAGES IA PRIORITAIRES" in text.upper() or "CAS D'USAGES" in text.upper():
                    current_section = "use_cases"
                    current_family = None  # Réinitialiser la famille au début de la section
                    logger.debug("Section 'Cas d'usage' détectée")
                    continue
                
                # --- Extraction des besoins ---
                if current_section == "needs":
                    if text.startswith("🔹"):
                        if current_need:
                            needs.append(current_need)
                        current_need = {"title": text.replace("🔹", "").strip(), "quotes": []}
                    elif ("•" in text or text.startswith("-")) and current_need:
                        # Retire la puce et espaces
                        clean_text = re.sub(r"^[•\-\s]+", "", text).strip()
                        # Extrait le contenu entre guillemets si présent
                        match = re.search(r"«(.*?)»", clean_text)
                        if match:
                            quote = match.group(1).strip()
                        else:
                            quote = clean_text
                        current_need["quotes"].append(quote)
                    print(f"current_need dans _try_extract_json : {current_need}")
                # --- Extraction des cas d'usage ---
                elif current_section == "use_cases":
                    # Détection d'un titre de famille
                    # Vérifier si c'est un titre de famille (style FamilyHeading ou format spécifique)
                    is_family_heading = False
                    try:
                        # Vérifier le style du paragraphe
                        style_name = para.style.name if para.style else None
                        if style_name == "FamilyHeading":
                            is_family_heading = True
                    except:
                        pass
                    
                    # Vérifier aussi si c'est "Autres cas d'usage" (titre de section sans famille)
                    if text == "Autres cas d'usage":
                        current_family = None
                        continue
                    
                    # Si c'est un titre de famille, mettre à jour current_family
                    if is_family_heading:
                        current_family = text.strip()
                        logger.debug(f"Titre de famille détecté: {current_family}")
                        continue
                    
                    # Nouveau cas d'usage (numéro suivi de titre)
                    if re.match(r"^\d+[\.\)]\s*", text):
                        # Sauvegarder le cas d'usage précédent
                        if current_use_case:
                            use_cases.append(current_use_case)
                        
                        # Extraire le titre
                        title = re.sub(r"^\d+[\.\)]\s*", "", text).strip()
                        current_use_case = {
                            "titre": title,
                            "description": "",
                            "famille": current_family  # Associer la famille courante
                        }
                    # Description du cas d'usage
                    elif (text.startswith("Description :") or text.startswith("Description:")) and current_use_case:
                        description = re.sub(r"^Description\s*:\s*", "", text, flags=re.IGNORECASE).strip()
                        current_use_case["description"] = description
                    # Autre texte pour le cas d'usage actuel
                    elif current_use_case:
                        if current_use_case["description"]:
                            current_use_case["description"] += " " + text
                        else:
                            current_use_case["description"] = text
            
            # Ajouter le dernier besoin et cas d'usage
            if current_need:
                needs.append(current_need)
            if current_use_case:
                use_cases.append(current_use_case)
            
            # Convertir les besoins au format attendu
            final_needs = []
            for need in needs:
                # Construire la description à partir des quotes et description
                description_parts = []
                if need.get("quotes"):
                    description_parts.extend(need["quotes"])
                if need.get("description"):
                    description_parts.append(need["description"])
                
                final_needs.append({
                    "titre": need.get("title", ""),
                    "description": " ".join(description_parts) if description_parts else ""
                })
            print(f"Final needs de l'extractor : {final_needs}")
            
            # Convertir les cas d'usage (la famille est déjà extraite depuis les titres)
            final_use_cases = []
            for uc in use_cases:
                titre = uc.get("titre", "")
                description = uc.get("description", "")
                famille = uc.get("famille")  # La famille a déjà été extraite depuis le titre
                
                final_use_cases.append({
                    "titre": titre,
                    "description": description,
                    "famille": famille
                })
            
            # Vérifier si on a extrait quelque chose
            if final_needs or final_use_cases:
                logger.info(f"✅ Extraction structurée réussie: {len(final_needs)} besoins, "
                           f"{len(final_use_cases)} cas d'usage")
                return {
                    "final_needs": final_needs,
                    "final_use_cases": final_use_cases
                }
            else:
                logger.debug("Aucune donnée structurée trouvée, utilisation de l'extraction LLM")
                return None
            
        except Exception as e:
            logger.warning(f"Erreur lors de l'extraction structurée: {e}", exc_info=True)
            return None
    
    def _extract_with_llm(self, word_path: str) -> Dict[str, List[Dict[str, Any]]]:
        """
        Extrait les données via LLM avec structured output.
        
        Args:
            word_path: Chemin vers le fichier Word
            
        Returns:
            Dict avec les données extraites
        """
        try:
            # Extraire le texte brut du Word
            doc = Document(word_path)
            word_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            if not word_text.strip():
                logger.warning("Document Word vide")
                return {
                    "final_needs": [],
                    "final_use_cases": []
                }
            
            # Préparer le prompt
            prompt = WORD_REPORT_EXTRACTION_PROMPT.format(word_text=word_text)
            
            # Appel à l'API avec structured output
            response = self.client.responses.parse(
                model=self.model,
                instructions="Tu es un expert en extraction de données structurées depuis des documents Word.",
                input=[
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                text_format=WordReportExtraction
            )
            
            # Extraire les données
            extracted_data = response.output_parsed.model_dump()
            
            logger.info(f"✅ Extraction LLM réussie: {len(extracted_data.get('final_needs', []))} besoins, "
                       f"{len(extracted_data.get('final_use_cases', []))} cas d'usage")
            
            return {
                "final_needs": extracted_data.get("final_needs", []),
                "final_use_cases": extracted_data.get("final_use_cases", [])
            }
            
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction LLM: {e}", exc_info=True)
            # Retourner une structure vide en cas d'erreur
            return {
                "final_needs": [],
                "final_use_cases": []
            }

